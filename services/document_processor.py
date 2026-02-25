"""Document processor for PDF, DOCX, XLSX, and TXT files."""

import io
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

from config.settings import settings


class DocumentProcessor:
    """Process various document formats and convert to text chunks."""

    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def process_file(self, file_content: bytes, filename: str) -> List[Dict[str, Any]]:
        """Process a file and return text chunks with metadata."""
        ext = Path(filename).suffix.lower()

        # Extract text based on file type
        if ext == ".pdf":
            text = self._extract_pdf(file_content)
        elif ext == ".docx":
            text = self._extract_docx(file_content)
        elif ext == ".xlsx":
            text = self._extract_xlsx(file_content)
        elif ext == ".txt" or ext == ".md":
            text = self._extract_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        # Split into chunks
        chunks = self._split_text(text)

        # Add metadata to each chunk
        results = []
        for i, chunk in enumerate(chunks):
            results.append(
                {
                    "content": chunk,
                    "metadata": {
                        "source": filename,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                    },
                }
            )

        return results

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            text = ""

            for page in reader.pages:
                text += page.extract_text() + "\n"

            return text
        except ImportError:
            raise ImportError("pypdf is required for PDF processing")

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from Word document."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            return text
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing")

    def _extract_xlsx(self, content: bytes) -> str:
        """Extract text from Excel file."""
        try:
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            text = ""

            for sheet in wb.sheetnames:
                ws = wb[sheet]
                text += f"\n## Sheet: {sheet}\n"

                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(
                        str(cell) if cell is not None else "" for cell in row
                    )
                    if row_text.strip():
                        text += row_text + "\n"

            return text
        except ImportError:
            raise ImportError("openpyxl is required for XLSX processing")

    def _extract_txt(self, content: bytes) -> str:
        """Extract text from plain text file."""
        return content.decode("utf-8", errors="ignore")

    def _split_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        try:
            from langchain_text_splitters import RecursiveCharacterTextSplitter

            splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                length_function=len,
            )

            return splitter.split_text(text)
        except ImportError:
            # Fallback to simple chunking
            return self._simple_chunking(text)

    def _simple_chunking(self, text: str) -> List[str]:
        """Simple chunking fallback."""
        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]
            chunks.append(chunk.strip())
            start = end - self.chunk_overlap

        return [c for c in chunks if c]


# Singleton instance
document_processor = DocumentProcessor()
